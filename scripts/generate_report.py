"""
Generate ProjectReport.docx for EC7205 Cloud Computing — Group 14
Event Ticket Booking Platform
"""
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "ProjectReport.docx")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin   = Cm(2.54)
section.right_margin  = Cm(2.54)

# ── Style helpers ─────────────────────────────────────────────────────────────
def add_heading(text, level=1, color="1F3864"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h

def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph(text, style="Normal")
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    for run in p.runs:
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(11)
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def set_cell_shading(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def style_table(tbl, header_fill="2E5090", header_text="FFFFFF"):
    tbl.style = "Table Grid"
    # Header row
    hdr = tbl.rows[0]
    for cell in hdr.cells:
        set_cell_shading(cell, header_fill)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(header_text)
                run.font.size = Pt(10)
    # Body rows
    for row in tbl.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(60)
run = title_p.add_run("UNIVERSITY OF RUHUNA")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.add_run("Faculty of Engineering\n").font.size = Pt(12)
run2 = sub_p.add_run("EC7205: CLOUD COMPUTING\n")
run2.bold = True
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
sub_p.add_run("Project Report — Assignment 2\n").font.size = Pt(12)
sub_p.add_run("Semester 7 — April 2026").font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

proj_p = doc.add_paragraph()
proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = proj_p.add_run("Event Ticket Booking Platform")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x2E, 0x50, 0x90)

doc.add_paragraph()
doc.add_paragraph()

team_p = doc.add_paragraph()
team_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
team_p.add_run("Group 14\n").font.size = Pt(12)
team_p.add_run("EG/2020/4044  –  Lakshani P.N.K\n").font.size = Pt(11)
team_p.add_run("EG/2020/4047  –  Lelwala L.G.S.R\n").font.size = Pt(11)
team_p.add_run("EG/2020/4054  –  Madhushani G.K.H.P\n").font.size = Pt(11)
team_p.add_run("EG/2020/4277  –  Welikumbura R.W.R.L").font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. Introduction")

add_body(
    "This report details the design, implementation, and deployment of a cloud-native "
    "Event Ticket Booking Platform developed as part of the EC7205 Cloud Computing module "
    "at the University of Ruhuna. The system enables users to browse events, reserve seats, "
    "and process payments, while administrators can manage event listings."
)
add_body(
    "The platform is built on a microservices architecture and deployed to Google Kubernetes "
    "Engine (GKE) using a fully automated CI/CD pipeline. It demonstrates core cloud computing "
    "principles including scalability, high availability, fault isolation, synchronous and "
    "asynchronous service communication, and modern DevOps practices."
)

add_heading("Key System Features", level=2)
add_bullet("User Authentication — JWT-based registration and login with role-based access control (Customer / Admin).")
add_bullet("Event Catalogue — Paginated, searchable event listing with seat availability display.")
add_bullet("Ticket Booking — Distributed seat locking via Redis SETNX to prevent double-booking under concurrent load.")
add_bullet("Payment Processing — Mock payment workflow that confirms bookings and triggers downstream notifications.")
add_bullet("Asynchronous Notifications — Event-driven notification consumer powered by RabbitMQ message broker.")
add_bullet("Admin Dashboard — Endpoint-level admin routes for creating and managing events.")
add_bullet("Horizontal Auto-scaling — HorizontalPodAutoscaler scales the booking service between 2–5 replicas based on CPU.")
add_bullet("HTTPS — Automatic TLS certificate provisioning via cert-manager and Let's Encrypt.")

add_heading("Technology Stack", level=2)

tbl = doc.add_table(rows=1, cols=2)
tbl.rows[0].cells[0].text = "Layer"
tbl.rows[0].cells[1].text = "Technology"

rows = [
    ("Backend Services",    ".NET 9, ASP.NET Core Web API, Entity Framework Core 9, FluentValidation"),
    ("Frontend",            "React 18, TypeScript, Vite 5, Tailwind CSS v4"),
    ("API Gateway",         "YARP Reverse Proxy 2.3, AspNetCoreRateLimit"),
    ("Database",            "PostgreSQL 16 (database-per-service pattern)"),
    ("Caching / Locking",   "Redis 7 (StackExchange.Redis, SETNX seat locks)"),
    ("Messaging",           "RabbitMQ 3 (RabbitMQ.Client 7, booking.created / payment.completed events)"),
    ("Orchestration",       "Google Kubernetes Engine (GKE Autopilot, us-central1)"),
    ("Container Registry",  "Google Artifact Registry"),
    ("Ingress / TLS",       "NGINX Ingress Controller, cert-manager, Let's Encrypt"),
    ("CI/CD",               "GitHub Actions, Workload Identity Federation (keyless auth)"),
    ("Local Dev",           "Docker Compose"),
]
for label, tech in rows:
    row = tbl.add_row()
    row.cells[0].text = label
    row.cells[1].text = tech
style_table(tbl)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. System Architecture")

add_body(
    "The application is composed of seven independent, containerized microservices deployed on a "
    "GKE Autopilot cluster. Each service owns its own PostgreSQL database (database-per-service "
    "pattern) to ensure complete data isolation and independent deployability. Communication is "
    "handled through a combination of synchronous REST calls via the API Gateway and asynchronous "
    "event publishing via RabbitMQ."
)

# Architecture diagram (text-based)
add_heading("Architecture Diagram", level=2)
add_body(
    "Figure 1 illustrates the high-level architecture of the platform:",
    bold=True
)
arch = doc.add_paragraph(style="Normal")
arch.paragraph_format.space_after = Pt(4)
arch_run = arch.add_run(
    "User / Browser\n"
    "      │  HTTPS\n"
    "      ▼\n"
    "┌─────────────────────────┐\n"
    "│   GCP Load Balancer     │\n"
    "└────────────┬────────────┘\n"
    "             │\n"
    "┌────────────▼────────────┐\n"
    "│  NGINX Ingress (K8s)   │\n"
    "└──────────┬──────┬───────┘\n"
    "  /api/*   │      │  /*\n"
    "           │      │\n"
    "┌──────────▼──┐  ┌▼────────────────┐\n"
    "│ API Gateway │  │  React Frontend  │\n"
    "│ (YARP+JWT)  │  │  (nginx:alpine)  │\n"
    "└──────┬──────┘  └─────────────────┘\n"
    "       │\n"
    "       ├─ /api/auth,/api/users  ─►  User Service    → PostgreSQL (users_db)\n"
    "       ├─ /api/events           ─►  Event Service   → PostgreSQL (events_db)\n"
    "       ├─ /api/bookings         ─►  Booking Service → PostgreSQL (bookings_db)\n"
    "       │                              (2–5 pods HPA)   → Redis (seat locks)\n"
    "       └─ /api/payments         ─►  Payment Service → PostgreSQL (payments_db)\n"
    "\n"
    "RabbitMQ (async events):\n"
    "  Booking Service  ──booking.created──►  Notification Service\n"
    "  Payment Service  ──payment.completed─► Notification Service\n"
    "\n"
    "CI/CD Pipeline:\n"
    "  GitHub Push → GitHub Actions → Artifact Registry → kubectl apply → GKE Rolling Update\n"
    "  (Workload Identity Federation — no GCP keys stored in GitHub)"
)
arch_run.font.name = "Courier New"
arch_run.font.size = Pt(8)

add_body("Figure 1: High-level system architecture.", italic=True)

add_heading("Services Breakdown", level=2)

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.rows[0].cells[0].text = "#"
tbl2.rows[0].cells[1].text = "Service"
tbl2.rows[0].cells[2].text = "Responsibility"

svc_rows = [
    ("1", "User Service",         "Registration, JWT login, ASP.NET Identity, role management (Admin / Customer)."),
    ("2", "Event Service",        "Event CRUD, seat management, paginated/filtered event listing."),
    ("3", "Booking Service",      "Seat reservation with Redis SETNX distributed lock; publishes booking.created to RabbitMQ. Runs 3 replicas with HPA."),
    ("4", "Payment Service",      "Mock payment processing (500 ms simulated delay); calls Booking Service to confirm; publishes payment.completed."),
    ("5", "Notification Service", "BackgroundService consuming booking.created and payment.completed queues; logs and simulates email dispatch."),
    ("6", "API Gateway",          "YARP reverse proxy; JWT Bearer authentication; IP rate limiting (100 req/min general, 20/min auth routes); routes /api/* to backend services."),
    ("7", "Frontend",             "React 18 + TypeScript SPA served by nginx:alpine; 7 pages; calls API Gateway via relative URLs at same Ingress host."),
]
for num, name, desc in svc_rows:
    row = tbl2.add_row()
    row.cells[0].text = num
    row.cells[1].text = name
    row.cells[2].text = desc
style_table(tbl2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. IMPLEMENTATION STEPS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. Implementation Steps")

add_heading("Step 1 — GCP & Cluster Setup", level=2)
add_body(
    "A GKE Autopilot cluster was provisioned in us-central1 using gcloud CLI. Autopilot was "
    "selected to eliminate manual node management while ensuring automatic bin-packing and resource "
    "efficiency within the $300 free-tier budget. An Artifact Registry Docker repository was created "
    "to store all container images. NGINX Ingress Controller and cert-manager were installed via Helm "
    "to handle external traffic and automatic HTTPS certificate provisioning from Let's Encrypt."
)

add_heading("Step 2 — Kubernetes Manifests", level=2)
add_body(
    "All infrastructure is defined as Kubernetes YAML manifests under the k8s/ directory. "
    "PostgreSQL instances are deployed as StatefulSets with persistent volume claims (standard-rwo, 5 Gi). "
    "Redis and RabbitMQ are also StatefulSets with persistent storage. All five microservices, the "
    "API gateway, and the frontend are Deployment resources with readiness/liveness probes on /health. "
    "Init containers ensure services wait for their database and messaging dependencies before starting. "
    "Non-sensitive configuration is stored in a ConfigMap; secrets (passwords, JWT key) are stored in "
    "a Kubernetes Secret, never in source control."
)

add_heading("Step 3 — CI/CD Pipeline (GitHub Actions)", level=2)
add_body(
    "A two-stage GitHub Actions workflow (.github/workflows/gke-deploy.yml) was created. On every "
    "push to main, the pipeline: (1) builds and validates the .NET solution and React app; "
    "(2) builds seven Docker images in parallel using docker/build-push-action and pushes both "
    ":sha and :latest tags to Artifact Registry; (3) authenticates to GCP via Workload Identity "
    "Federation (keyless — no SA JSON keys stored in GitHub); (4) gets GKE credentials, substitutes "
    "image tags in manifests using sed, creates/updates Kubernetes secrets from GitHub secrets, "
    "applies all manifests, and waits for rolling deployments to complete."
)

add_heading("Step 4 — Microservices (ASP.NET Core)", level=2)
add_body(
    "Each microservice follows the database-per-service pattern: it owns a dedicated PostgreSQL "
    "database and runs EF Core migrations on startup (EnsureCreated). The Booking Service "
    "implements distributed seat locking using Redis SETNX (key: seat-lock:{eventId}:{seatId}, "
    "10-second TTL) to prevent race conditions under concurrent load. RabbitMQ communication uses "
    "the RabbitMQ.Client v7 async API. The API Gateway uses YARP with JWT validation middleware "
    "and AspNetCoreRateLimit for API-level throttling."
)

add_heading("Step 5 — Frontend (React + Vite)", level=2)
add_body(
    "The React 18 frontend uses TypeScript, Tailwind CSS v4, TanStack React Query, and React "
    "Router v6. In production, VITE_API_URL is set to an empty string so axios uses relative "
    "URLs (e.g., /api/auth/login). The NGINX Ingress routes /api/* to the API Gateway and /* to "
    "the frontend, enabling both to share the same HTTPS host (https://<IP>.nip.io). The frontend "
    "nginx.conf uses try_files $uri $uri/ /index.html for SPA client-side routing."
)

add_heading("Step 6 — Security", level=2)
add_body(
    "JWT Bearer tokens are validated at the API Gateway before requests reach any backend service. "
    "The gateway applies IP rate limiting (100 req/min general, 20/min for auth endpoints) to "
    "mitigate brute-force attacks. HTTPS is enforced via cert-manager with automatic Let's Encrypt "
    "certificate renewal. Passwords are hashed with ASP.NET Identity (PBKDF2). Kubernetes Secrets "
    "store sensitive values and are never committed to source control. CORS is restricted to "
    "known origins."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. CHALLENGES FACED & SOLUTIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading("4. Challenges Faced & Solutions")

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.rows[0].cells[0].text = "Challenge"
tbl3.rows[0].cells[1].text = "Problem Description"
tbl3.rows[0].cells[2].text = "Solution Implemented"

challenges = [
    (
        "Booking Service Cluster in Gateway",
        "The YARP API Gateway was configured to route bookings through an nginx-lb container (needed for local Docker Compose scaling). In GKE, this container does not exist — Kubernetes Service natively load-balances across pods.",
        "Added appsettings.Production.json to the gateway project that overrides the booking-service cluster address to http://booking-service:8080. In GKE, ASP.NET Core automatically loads this file in the Production environment, pointing YARP to the Kubernetes Service instead of nginx-lb."
    ),
    (
        "Frontend API URL",
        "The React frontend baked the API URL (http://localhost:5000) into the bundle at build time via Vite's import.meta.env. In GKE, the API is served at https://<IP>.nip.io — an IP not known until after cluster creation.",
        "Changed the axios base URL fallback from || (falsy check) to ?? (nullish coalescing). Setting VITE_API_URL='' (empty string) in the Docker build-arg causes axios to use relative URLs. The NGINX Ingress then routes /api/* to the gateway and /* to the frontend from the same host, so no hardcoded IP is needed."
    ),
    (
        "NGINX Ingress path conflicts",
        "NGINX Ingress was routing both / and /api/ paths but serving the frontend for all API calls because the more specific /api prefix rule was being shadowed.",
        "Added nginx.ingress.kubernetes.io/use-regex: 'true' annotation and ensured /api (Prefix) is declared before / (Prefix) in the Ingress rules. NGINX evaluates longer prefix matches first, correctly sending /api/* to the gateway and /* to the frontend."
    ),
    (
        "GKE Autopilot resource minimums",
        "GKE Autopilot enforces a minimum of 250m CPU and 512Mi memory per container. Initial manifests with 100m/128Mi requests were rejected.",
        "Updated all container resource requests to 250m CPU / 512Mi memory and limits to 500m / 1Gi. This is Autopilot's documented minimum and ensures pods are admitted by the scheduler."
    ),
]
for title, problem, solution in challenges:
    row = tbl3.add_row()
    row.cells[0].text = title
    row.cells[1].text = problem
    row.cells[2].text = solution
style_table(tbl3)

doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. LESSONS LEARNED
# ══════════════════════════════════════════════════════════════════════════════
add_heading("5. Lessons Learned")

lessons = [
    (
        "Infrastructure as Code is Essential",
        "Defining all Kubernetes resources in YAML files managed via Git provided a single source of "
        "truth. Every cluster change was traceable, deployments were reproducible across environments, "
        "and rollbacks required only a git revert followed by re-applying the manifests."
    ),
    (
        "Decoupling with Event-Driven Architecture",
        "Using RabbitMQ to communicate between the Booking, Payment, and Notification services "
        "proved highly effective. A failure or slowdown in the Notification Service had zero impact "
        "on the core booking flow, demonstrating the resilience benefit of asynchronous decoupling."
    ),
    (
        "Kubernetes Native Load Balancing Eliminates Custom Proxies",
        "Local development required a dedicated NGINX container to load-balance across multiple "
        "booking-service replicas. In Kubernetes, the ClusterIP Service provides this transparently. "
        "The database-per-service pattern and headless services for StatefulSets gave each component "
        "stable, predictable DNS names regardless of pod restarts."
    ),
    (
        "Keyless CI/CD Authentication is More Secure than Key Files",
        "Workload Identity Federation allowed GitHub Actions to authenticate to GCP without storing "
        "any service account JSON key in GitHub Secrets. This approach eliminates a significant "
        "attack surface (credential leakage from secret exposure) and follows zero-trust principles."
    ),
    (
        "Observability Must Be Designed In, Not Added Later",
        "The /health endpoint on every service enabled Kubernetes readiness and liveness probes that "
        "prevented unhealthy pods from receiving traffic. kubectl logs and kubectl describe pod were "
        "critical for diagnosing init container failures, missing secrets, and database connection "
        "timing issues during the initial deployment."
    ),
]
for title, body in lessons:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"\u2022  {title}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(body)
    r2.font.size = Pt(11)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
